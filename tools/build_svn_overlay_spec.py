from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"D:\work\lsl\9.upper\1.code\27.SVN_tools\SVN_悬浮助手_设计规格文档.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "222222"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "C9D1D9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        tag = "w:" + side
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def set_run_font(run, size=None, bold=None, color=None, name="Microsoft YaHei UI"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", size=10.5, color=INK, bold=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 1:
        size, color, before, after = 16, BLUE, 18, 10
    elif level == 2:
        size, color, before, after = 13, BLUE, 14, 7
    else:
        size, color, before, after = 12, DARK_BLUE, 10, 5
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    set_cell_border(cell, "D0D7DE")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=10, bold=True, color=DARK_BLUE)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, value in enumerate(row):
            cell = cells[ci]
            set_cell_shading(cell, WHITE if ri % 2 == 0 else "FAFBFC")
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(str(value))
            set_run_font(run, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=9.5, color="333333", name="Consolas")
    return p


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei UI"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei UI"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SVN Floating Assistant 设计规格")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SVN 悬浮助手 · V1.0")
    set_run_font(run, size=9, color=MUTED)


def build():
    doc = Document()
    configure_doc(doc)

    add_para(doc, "SVN Floating Assistant", size=26, color=DARK_BLUE, bold=True, after=3)
    add_para(doc, "Windows 桌面 SVN 信息悬浮助手设计规格文档", size=14, color=MUTED, after=12)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["版本", "V1.0"],
            ["目标平台", "Windows 10/11"],
            ["技术栈", "C# / .NET 8 / WPF / TortoiseSVN / svn.exe"],
            ["定位", "轻量级 SVN 信息展示与快捷入口，不替代 SVN 客户端"],
            ["当前代码路径", r"D:\work\lsl\9.upper\1.code\27.SVN_tools\SvnFloatingAssistant"],
        ],
        [1800, 7560],
    )

    add_note(
        doc,
        "本文档按当前项目真实实现整理：圆形悬浮球、展开信息面板、Directory Opus 路径识别、TortoiseSVN/SubWCRev 兼容模式、svn.exe 完整模式，以及 TortoiseProc 快捷入口。"
    )

    add_heading(doc, "1. 项目定位", 1)
    add_para(
        doc,
        "SVN Floating Assistant 是一个常驻桌面的轻量级 SVN 信息助手。它关注用户当前正在浏览的目录，自动判断目录是否属于 SVN 工作副本，并以悬浮球和信息面板的方式展示关键状态。"
    )
    add_bullet(doc, "信息展示：当前目录状态、最后变更作者、日期、本地修改概要、最近提交日志。")
    add_bullet(doc, "快捷入口：打开目录、TortoiseSVN 日志、提交、更新。")
    add_bullet(doc, "低侵入：不主动扫描磁盘，不替代 TortoiseSVN，不改变用户已有 SVN 工作流。")
    add_bullet(doc, "可扩展：后续可扩展 Git、Keil 工程识别、Hex 版本读取、固件版本展示。")

    add_heading(doc, "2. 用户场景", 1)
    add_table(
        doc,
        ["场景", "用户期望", "工具响应"],
        [
            ["浏览 SVN 工程目录", "快速知道当前是否有修改", "悬浮球变为绿色/黄色/红色，展开面板显示状态概要"],
            ["切换到非 SVN 目录", "不被无关信息打扰", "显示非 SVN 状态或保持低信息量提示"],
            ["需要看最近提交", "不用进入完整日志窗口", "展开面板按日期显示最近日志"],
            ["需要提交或更新", "继续使用熟悉的 TortoiseSVN", "调用 TortoiseProc.exe 打开原生对话框"],
            ["使用 Directory Opus", "像资源管理器一样被识别", "从 Directory Opus 当前窗口标题读取路径并验证目录存在"],
        ],
        [2200, 3300, 3860],
    )

    add_heading(doc, "3. 当前实现架构", 1)
    add_table(
        doc,
        ["模块", "职责", "关键文件"],
        [
            ["UI", "悬浮球、展开面板、右键菜单、状态颜色", "MainWindow.xaml / MainWindow.xaml.cs"],
            ["ViewModel", "将 SVN 状态转换为中文展示文案和日志列表", "ViewModels/MainViewModel.cs"],
            ["路径监听", "识别前台 Explorer 或 Directory Opus 当前目录", "Services/ExplorerPathMonitor.cs"],
            ["SVN 服务", "执行 svn.exe 或 SubWCRev.exe，解析 info/status/log", "Services/SvnService.cs"],
            ["缓存", "缓存 info/status/log，减少重复调用", "Services/SvnCache.cs"],
            ["快捷入口", "调用 TortoiseProc.exe 打开日志、提交、更新", "Services/TortoiseSvnLauncher.cs"],
            ["配置", "JSON 设置文件，支持工具路径和刷新参数", "Services/AppSettings.cs"],
        ],
        [1800, 4260, 3300],
    )

    add_heading(doc, "4. UI 设计", 1)
    add_heading(doc, "4.1 悬浮球", 2)
    add_para(doc, "悬浮球默认显示在屏幕右侧，作为低干扰状态指示器。")
    add_table(
        doc,
        ["颜色", "短文案", "含义"],
        [
            ["灰色", "OFF", "未识别到 SVN 工作副本或等待路径"],
            ["绿色", "正常", "SVN 工作副本无本地修改"],
            ["黄色", "改动", "存在本地修改、未加入文件或本地缺失"],
            ["红色", "冲突", "存在冲突文件"],
            ["蓝色", "...", "正在刷新"],
        ],
        [1600, 1600, 6160],
    )
    add_heading(doc, "4.2 展开面板", 2)
    add_bullet(doc, "标题区显示项目名和当前路径。")
    add_bullet(doc, "状态区显示工作副本状态、最后变更作者、最后变更日期。")
    add_bullet(doc, "修改区使用中文描述，例如“已修改 5 个，未加入版本库 2 个”。")
    add_bullet(doc, "日志区按日期和作者展示最近提交，不突出 revision 编号。")
    add_bullet(doc, "底部按钮提供刷新、更新、提交、日志四个快捷操作。")
    add_note(doc, "已修复展开/收起时窗口位置漂移：切换面板宽度时保持右边缘位置不变。")

    add_heading(doc, "5. 路径监听设计", 1)
    add_table(
        doc,
        ["来源", "识别方式", "说明"],
        [
            ["Windows Explorer", "Shell.Application COM + 当前窗口 HWND 匹配", "读取 LocationURL 并转换为本地路径"],
            ["Directory Opus", "检测前台进程并读取窗口标题", "标题为有效目录时作为当前路径"],
            ["其他程序", "忽略", "避免误把编辑器或浏览器标题当路径"],
        ],
        [2200, 3560, 3600],
    )
    add_bullet(doc, "默认 500ms 检测一次前台路径。")
    add_bullet(doc, "路径变化后等待 300~500ms 再刷新，避免快速切换目录时频繁调用 SVN。")
    add_bullet(doc, "路径未变化时不刷新。")

    add_heading(doc, "6. SVN 数据获取模式", 1)
    add_heading(doc, "6.1 完整模式：svn.exe", 2)
    add_para(doc, "当检测到 svn.exe 时，工具进入完整模式，支持精确状态统计和最近日志。")
    add_table(
        doc,
        ["命令", "超时", "用途"],
        [
            ["svn info", "1 秒", "判断是否为 SVN 工作副本，读取 Revision、Repository URL、Author、Date"],
            ["svn status", "3 秒", "统计已修改、已新增、已删除、未加入版本库、本地缺失、冲突数量"],
            ["svn log -l 5 --xml", "3 秒", "读取最近 5 条提交日志，按日期和作者展示"],
        ],
        [3500, 1300, 4560],
    )

    add_heading(doc, "6.2 兼容模式：SubWCRev.exe", 2)
    add_para(doc, "当没有 svn.exe，但存在 TortoiseSVN 自带的 SubWCRev.exe 时，工具进入兼容模式。")
    add_bullet(doc, "可以判断目录是否为 SVN 工作副本。")
    add_bullet(doc, "可以读取 Last committed revision 和 Updated revision。")
    add_bullet(doc, "可以粗略判断是否存在本地修改或未版本控制项。")
    add_bullet(doc, "不读取完整日志，不做精确 M/A/D/? 分类统计。")
    add_note(doc, "兼容模式的目标是让只安装 TortoiseSVN 图形客户端的机器也能显示基础状态；完整日志和精确统计仍建议安装 svn.exe 命令行组件。")

    add_heading(doc, "7. 缓存与性能策略", 1)
    add_table(
        doc,
        ["数据", "缓存时间", "刷新条件"],
        [
            ["Info", "60 秒", "路径变化或缓存过期"],
            ["Status", "10 秒", "路径变化、手动刷新或缓存过期"],
            ["Log", "60 秒", "面板展开且缓存过期"],
        ],
        [1800, 1800, 5760],
    )
    add_bullet(doc, "所有 SVN 命令均在后台任务执行，禁止阻塞 UI 线程。")
    add_bullet(doc, "所有进程调用都有超时控制，超时后显示“SVN响应慢”。")
    add_bullet(doc, "不递归扫描磁盘，不使用 Directory.GetFiles(*) 做全盘枚举。")
    add_bullet(doc, "日志按需加载：悬浮球状态不加载日志，展开面板后才加载。")

    add_heading(doc, "8. 快捷入口", 1)
    add_table(
        doc,
        ["入口", "调用方式", "效果"],
        [
            ["打开目录", "Process.Start(path)", "打开当前目录"],
            ["日志", "TortoiseProc.exe /command:log", "打开 TortoiseSVN 日志窗口"],
            ["提交", "TortoiseProc.exe /command:commit", "打开 TortoiseSVN 提交窗口"],
            ["更新", "TortoiseProc.exe /command:update", "打开 TortoiseSVN 更新窗口"],
        ],
        [1800, 3600, 3960],
    )
    add_note(doc, "快捷入口继续交给 TortoiseSVN 原生窗口完成，避免工具承担 Merge、Resolve、Diff 等 SVN 客户端职责。")

    add_heading(doc, "9. 设置文件", 1)
    add_para(doc, r"配置文件路径：%APPDATA%\SvnFloatingAssistant\settings.json")
    for line in [
        "{",
        '  "AutoRefresh": true,',
        '  "ExplorerPollMilliseconds": 500,',
        '  "DebounceMilliseconds": 400,',
        '  "BubbleSize": 72,',
        '  "Opacity": 0.95,',
        '  "DarkMode": false,',
        '  "SvnPath": null,',
        '  "SubWCRevPath": null,',
        '  "TortoiseSvnProcPath": null',
        "}",
    ]:
        add_code(doc, line)

    add_heading(doc, "10. 已完成能力", 1)
    for item in [
        "WPF .NET 8 工程骨架。",
        "悬浮球和展开信息面板。",
        "Explorer 当前路径监听。",
        "Directory Opus 当前路径识别。",
        "svn.exe 完整模式：info/status/log。",
        "SubWCRev.exe 兼容模式。",
        "TortoiseProc.exe 日志、提交、更新快捷入口。",
        "SVN 状态中文化展示。",
        "最近日志按日期和作者展示。",
        "展开/收起位置漂移修复。",
        "bin、obj SVN 忽略规则。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. 后续规划", 1)
    add_table(
        doc,
        ["版本", "功能", "说明"],
        [
            ["V1.0", "基础信息展示", "悬浮球、状态识别、最近日志、快捷入口"],
            ["V1.1", "修改文件列表", "展开后显示已修改、新增、删除、冲突文件明细"],
            ["V1.2", "电源研发扩展", "Keil/UVProj/Hex 识别、Firmware Version、Boot Version、Build 时间"],
            ["V1.3", "多版本控制支持", "Git/SVN/Mercurial 自动识别"],
        ],
        [1400, 2600, 5360],
    )

    add_heading(doc, "12. 风险与约束", 1)
    add_bullet(doc, "公司 SVN 服务器权限完全由服务器控制，工具不会绕过权限。")
    add_bullet(doc, "自动提交不是目标功能，提交仍建议通过 TortoiseSVN 原生窗口手动确认。")
    add_bullet(doc, "中文提交日志通过命令行提交时必须使用 UTF-8 文件和 --encoding UTF-8，避免 PowerShell 参数编码乱码。")
    add_bullet(doc, "如果服务器未启用 pre-revprop-change hook，历史提交日志无法由普通用户修改。")
    add_bullet(doc, "Directory Opus 标题格式若被用户自定义，可能需要额外路径解析规则。")

    doc.add_section(WD_SECTION.CONTINUOUS)
    add_para(doc, "文档结束", size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=12)
    doc.save(OUT)


if __name__ == "__main__":
    build()
